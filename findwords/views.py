#-*- coding: UTF-8 -*-
from django.shortcuts import render
from django.template.loader import get_template
from django.http import HttpResponse, JsonResponse
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from docx import Document
import os,sys,datetime
import logging
import json

logging = logging.getLogger('reso_logger')
# logging.info("iiiiiiiiiiiiiiiiiiiiiiiiii")
# Create your views here.

# 解决datetime、date格式数据无法json序列化问题
class DateEncoder(json.JSONEncoder):  
    def default(self, obj):  
        if isinstance(obj, datetime.datetime):  
            return obj.strftime('%Y-%m-%d %H:%M:%S')
        elif isinstance(obj, datetime.date):
            return obj.strftime("%Y-%m-%d")  
        else:  
            return json.JSONEncoder.default(self, obj) 

def index(request):
    return HttpResponse("Hello, world. You're at the findwords index.")

def search_string(filename,string):
    #打开文档
    document = Document(filename)
    # document = Document(r'C:\Users\Cheng\Desktop\kword\words\wind.docx')
    print filename
    #读取每段资料
    l = [ paragraph.text for paragraph in document.paragraphs];
    # l = [ paragraph.text.encode('gb2312') for paragraph in document.paragraphs];
    #输出并观察结果，也可以通过其他手段处理文本即可
    fileword = []
    for i in l:
        i=i.strip()
        # print i
        if i.find(string)!=-1:
            changetime = datetime.datetime.fromtimestamp(os.path.getmtime(filename)).strftime('%Y-%m-%d %H:%M:%S')
            logging.info(changetime)
            # print filename, i
            fword = filename+">>>>>"+changetime+">>>>>"+i
            fileword.append(fword)
    # logging.info(fileword)
    return fileword

#遍历该目录下的所有文件，返回‘目录+文件名’列表
def get_process_files(root_dir):
    """process all files in directory"""
    cur_dir=os.path.abspath(root_dir)
    file_list=os.listdir(cur_dir)
    logging.info(file_list)
    process_list=[]
    for file in file_list:
        fullfile=cur_dir+"\\"+file
        if os.path.isfile(fullfile):
            process_list.append(fullfile)
        elif os.path.isdir(fullfile):
            dir_extra_list=get_process_files(fullfile)
            if len(dir_extra_list)!=0:
                for x in dir_extra_list:
                    process_list.append(x)
    # print process_list
    return process_list

def count_files(root_dir,string):
    process_list=get_process_files(root_dir)
    logging.info(process_list)
    f_result = []
    for files in process_list:
        f_result.append(search_string(files, string))
    return f_result

def findwords(requset):
    return render(requset, 'findwords/search.html')

def result(requset):
    logging.info("ffffffffffffffffffffffffff")
    # return render(requset, "findwords/search.html")
    try:
        if (requset.GET['word']):
            word = requset.GET['word']
            logging.info(word)
            root_dir="..\\words" #目录
            string = word #要搜索的字符串
            # logging.info("dddddddddddddddddddddddddd")
            try:
                # logging.info("yyyyyyyyyyyyyyyyyyyyyyy")
                f_result = count_files(root_dir,string)
                logging.info(f_result)
                return render(requset, 'findwords/f_result.html', {"f_result": f_result})
            except:
                pass
            # return render(requset, 'findwords/search.html', locals())
    except:
        pass

def searchwords(request):
    # json_data = [{
    #         "fileId": "1",
    #         "filename": "倚天屠龙记",
    #         "checktime": "2018-10-11 08:00",
    #         "contents": "张无忌朗声道：“晚辈学过贵派的一些七伤拳法，倘若练得不对，请崆峒派各位前辈切莫见笑。”各派人众听了，尽皆诧异：“这小子原来连崆峒派的七伤拳也会，那是从何处学来啊？”只听他朗声念道：“五行之气调阴阳，损心伤肺摧肝肠，藏离精失意恍惚，三焦齐逆兮魂魄飞扬！”别派各人听到，那也罢了。崆峒五老听到他高吟这四句似歌非歌、似诗非诗的拳诀，却无不凛然心惊。这正是七伤拳的总诀，乃崆峒派的不传之秘，这少年如何知道？他们一时之间，怎想得到谢逊将七伤拳谱抢去后，传了给他。"
    #     }]
    return render(request, 'findwords/findwords.html', locals())

@csrf_exempt
def searesult(request):
    logging.info("ffffffffffffffffffffffffff")
    # json_data = [{
    #         "fileId": "1",
    #         "filename": "倚天屠龙记",
    #         "checktime": "2018-10-11 08:00",
    #         "contents": "张无忌朗声道：“晚辈学过贵派的一些七伤拳法，倘若练得不对，请崆峒派各位前辈切莫见笑。”各派人众听了，尽皆诧异：“这小子原来连崆峒派的七伤拳也会，那是从何处学来啊？”只听他朗声念道：“五行之气调阴阳，损心伤肺摧肝肠，藏离精失意恍惚，三焦齐逆兮魂魄飞扬！”别派各人听到，那也罢了。崆峒五老听到他高吟这四句似歌非歌、似诗非诗的拳诀，却无不凛然心惊。这正是七伤拳的总诀，乃崆峒派的不传之秘，这少年如何知道？他们一时之间，怎想得到谢逊将七伤拳谱抢去后，传了给他。"
    #     }]
    if request.method == 'POST':
        try:
            word = request.POST.get('keyword')
            logging.info(word)
            root_dir="..\\words" #目录
            string = word #要搜索的字符串
            try:
                f_result = count_files(root_dir,string)
                logging.info(f_result)
                f_result_list = []
                for result in f_result:
                    for res in result:
                        result_list = res.split('>>>>>')
                        result_dict = {"filename":result_list[0],"checktime":result_list[1],"contents":result_list[2]}
                        # logging.info(result_dict)
                        f_result_list.append(result_dict)
                logging.info(f_result_list)
                json_data = json.dumps(f_result_list, cls=DateEncoder, ensure_ascii=False)
                return HttpResponse(json_data, content_type="application/json")
            except:
                return render(request, 'findwords/findwords.html', locals())
        except:
            return render(request, 'findwords/findwords.html', locals())
    else:
        return render(request, 'findwords/findwords.html', locals())